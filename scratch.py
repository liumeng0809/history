# -*- coding: utf-8 -*-
"""通用朝代历史科普生成器 · 兴衰阶段主干 × 事件驱动支线

把任意朝代/政权讲透，对历史小白友好。
主干：让 AI 先把该朝代按兴衰分成 N 个阶段（不依赖手写阶段表，通吃所有朝代）。
支线：每个阶段内按时间顺序列出"完整历史事件"，一个事件只出现一次，扩写时集中讲透（起因/经过/结果/影响）。
输出：每个阶段一个 Word（{朝代}_1_{阶段名}.docx …）。

流程：背景总览（1次）→ 兴衰阶段划分（1次）→ 逐阶段事件大纲（每阶段1次）→ 串行扩写每事件 → 按阶段落盘 Word。
扩写沿用已验证约束：史实最高原则不编造、口语化、讲来龙去脉、名词解释、串行衔接、严禁前情回顾。

用法：只改下面 TOPIC 一行，填你想讲的朝代/政权名（如"唐朝历史""宋朝历史""罗马帝国"），运行即可。
"""
import re
from openai import OpenAI
import httpx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ================= 配置区域 =================
API_KEY = "sk-Qaes2KCUDHr67GZoif13ySqsuFsD7NYWAnLoVcNcAlv3mcXW"
MODEL_NAME = "DeepSeek-V4-Flash"
TOPIC = "东晋国历史"   # ← 只改这里：填你想讲的朝代/政权名
# ============================================

client = OpenAI(
    api_key=API_KEY,
    base_url="http://aiserver.hisi.huawei.com/v1",
    # 绕过 Windows 注册表里的 IE 系统代理（华为 proxycn2:8080），否则 504
    http_client=httpx.Client(trust_env=False)
)

_BOLD_RE = re.compile(r'\*\*(.+?)\*\*')


def add_runs_with_bold(paragraph, text):
    """把含 **加粗** 的文本拆成多个 run，加粗部分 bold=True；未配对的 ** 直接清除。"""
    if text.count('**') >= 2:
        pos = 0
        for m in _BOLD_RE.finditer(text):
            if m.start() > pos:
                paragraph.add_run(text[pos:m.start()])
            run = paragraph.add_run(m.group(1))
            run.bold = True
            pos = m.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])
    else:
        paragraph.add_run(text.replace('**', ''))


def call_ai(prompt, max_tokens=8000):
    """调用大模型的通用函数。max_tokens 可按需调高，防止长正文被截断。"""
    print("  思考中...", end="", flush=True)
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=max_tokens
    )
    print("完成!", flush=True)
    return response.choices[0].message.content


# ---- 第一步：朝代背景总览（全局事实基线）----
def generate_overview(topic):
    prompt = f"""
你是历史顾问。请为【{topic}】写一份背景总览，作为后续分阶段讲解的事实基线。
只输出事实，不要议论，按下面5个小标题逐段输出，每段3-5句话，必须含具体年份/地名/人名：
### 基本盘（起止时间、都城、开国与亡国之君、疆域大概）
### 制度骨架（核心政治/军事/经济制度的来龙去脉）
### 贯穿全程的主线矛盾（如皇权与文官、中央与地方、内地与边患、财政与开支等）
### 兴衰关键转折（哪几个节点决定走向，按时间列出）
### 对后世的影响
"""
    print(f"\n第一步：生成【{topic}】背景总览...")
    return call_ai(prompt)


# ---- 第二步：兴衰阶段划分（让 AI 按该朝代兴衰分阶段，通吃所有朝代）----
def generate_stages(topic, overview):
    """让 AI 把【topic】按兴衰分成若干阶段，返回 [(阶段序号, 阶段名, 年代, 核心说明), ...]。"""
    prompt = f"""
你是历史顾问。请为【{topic}】规划讲解主干——按"兴衰阶段"把它从头到尾分成若干阶段。

【背景总览】（事实基线，不得矛盾）：
{overview}

要求：
1. 按该朝代/政权的兴衰节奏分阶段，通常 5-8 个阶段为宜（太细则碎，太粗则讲不透）。
2. 每个阶段要能独立成章，对应一个相对完整的"兴"或"衰"的段落，不能把同一阶段的事件拆到两处。
3. 严格按下面格式输出，每行一个阶段，行内用竖线｜分隔四段，行首不要加序号符号：
   阶段序号｜阶段名（4-8字）｜年代范围｜本阶段核心说明（一句话，含关键事件）
   例如：
   1｜开国奠基｜1368-1402｜朱元璋建国、废丞相、洪武之治、靖难之役
   2｜永宣盛世｜1403-1435｜永乐迁都、郑和下西洋、仁宣之治
4. 阶段序号从1开始连续递增，年代范围覆盖该朝代从建立到灭亡的完整时间。
5. 只输出阶段列表，不要输出标题、解释或多余文字。
"""
    print(f"\n第二步：为【{topic}】划分兴衰阶段...")
    raw = call_ai(prompt)
    stages = []
    for ln in raw.splitlines():
        ln = ln.strip()  # 注意：不 lstrip 序号，序号是要解析的字段
        if not ln or '｜' not in ln:
            continue
        parts = ln.split('｜', 3)
        if len(parts) != 4:
            continue
        no, name, years, summary = [p.strip() for p in parts]
        if no.isdigit():
            stages.append((int(no), name, years, summary))
    # 按序号排序并重排为连续（防止 AI 跳号）
    stages.sort(key=lambda s: s[0])
    return [(i + 1, s[1], s[2], s[3]) for i, s in enumerate(stages)]


# ---- 第三步：单阶段大纲（事件驱动，一个事件只出现一次，集中讲透）----
def generate_stage_outline(topic, overview, stage_no, stage_name, stage_years, stage_summary, prior_stage_events=None):
    """为一个阶段生成大纲：该阶段有哪些必须讲到的完整历史事件，按时间顺序列出。
    每条 = 一个事件（含其起因/经过/结果/影响），后续扩写时集中讲透，不跨条重复。
    返回 [(event, point, year_start), ...]，year_start 为事件起始年份（int），供代码按时间排序。
    prior_stage_events: 前面阶段已列过的事件 [(事件名, 年份), ...]，用于跨阶段去重（防止文件间内容重复）。"""
    prior_hint = ""
    if prior_stage_events:
        prior_list = "\n".join(f"- {e}（{y}年）" for e, y in prior_stage_events)
        prior_hint = f"""
【前面阶段已经列过的事件——跨阶段不重复，重要】下面这些事件在前面阶段已作为大纲条目讲过，本阶段【不得再列】：
{prior_list}
本阶段只列前面阶段没讲过的新事件；若某事件前面讲过，本阶段不再单列，必要时靠后续"沿用前述"一句带过即可。
"""
    prompt = f"""
你是一位严谨的历史学家，面向完全不懂历史的读者做科普。
现在要为【{topic}】的第{stage_no}阶段【{stage_name}（{stage_years}）】生成讲解大纲。

【{topic}】背景总览（全局事实基线，不得矛盾）：
{overview}

本阶段核心：{stage_summary}
{prior_hint}
要求：
1. 列出本阶段必须讲到的【完整历史事件】，按时间顺序排列，通常 6-8 个事件为宜（不要超过 8 个，否则单阶段篇幅过大易截断）。
2. 【严格守本阶段年代范围——重要】本阶段年代范围是【{stage_years}】。只能列起始年份落在这个范围内的（含两端）事件；
   不得把发生在本阶段【之后】的事件塞进来（如本阶段到322年结束，就不得列323、324年的事件，那是下一阶段的内容）。
3. 【以时间线为准——重要】每个事件是时间线上的一个独立节点，按发生时间严格递进，沿时间线往前讲。
   同一时间点发生的同一件事（起因/经过/结果/影响）归一条；
   但不同时间点发生的分别独立成条，不要合并（如"王敦322年第一次起兵""晋明帝323年继位""王敦324年第二次起兵"是三个独立节点，各自成条，按年份排列，不要合并成"王敦之乱"一条）。
   【例外——同一制度/政策只算一条】同一项制度、政策或机制的"初创→铺开→反复推行"（如侨寄法/侨置郡县在317年提出、318年大规模铺开、后来多次土断），即使跨越多个时间点，也只算一个事件，合并为一条，事件名用制度名，起始年份取该制度最早出现之年，要点概述里写清它从提出到落实做了什么；不得拆成"初创""铺开""再推行"多条。
4. 事件之间不要重复：一个事件只在一条里出现，后面不再单列。
5. 既要有政治军事大事，也要兼顾本阶段重要的经济、社会、对外事件，保证小白能看到完整图景。
6. 每条大纲只写"具体事件名 + 起始年份 + 一句要点概述"，不要写正文。
   格式严格如下，每条一行，行首不要加序号符号，行内用竖线｜分隔三段：
   王敦第一次起兵｜322｜322年王敦以清君侧为名从武昌起兵攻入建康，逼死晋元帝
   晋明帝继位｜323｜323年太子司马绍继位，暗中蓄力对抗王敦
   王敦第二次起兵｜324｜324年王敦病重时再次起兵，被晋明帝讨平
   说明：第二段是该事件的【起始年份】（事件开始的那一年，四位数字，如322），用于排序；第三段是要点概述。
7. 事件名必须真实、具体，要点须含年份/关键人物，起始年份必须真实准确（必须是该事件开始之年，不是结束之年或中间之年），不得虚构。
8. 事件之间不要重复：一个事件只在一条里出现，后面不再单列。【已讲制度不重立】凡【本阶段已经讲过的事件】里已列出的制度/政策（如侨寄法、土断、侨置郡县），不得再以其后续落实、推广、反复推行为名另立条目——那是同一制度的延续，不是新事件。
"""
    print(f"\n  生成【阶段{stage_no}：{stage_name}】大纲...")
    raw = call_ai(prompt)
    items = []
    for ln in raw.splitlines():
        ln = ln.strip().lstrip('0123456789.-* ')
        if not ln or '｜' not in ln:
            continue
        parts = ln.split('｜', 2)
        if len(parts) != 3:
            continue
        event, year_str, point = [p.strip() for p in parts]
        # 提取起始年份：取 year_str 中第一个连续 3-4 位数字
        m = re.search(r'(\d{3,4})', year_str)
        if event and m:
            year_start = int(m.group(1))
            items.append((event, point, year_start))
    # 按起始年份升序排序（不信任 AI 的排列顺序，防止时间倒错）
    items.sort(key=lambda x: x[2])
    return items


def _normalize_event_name(name):
    """归一化事件名用于跨阶段去重：去掉"第X次/首次/起兵/之乱"等修饰与空格，取核心词组。"""
    s = name
    for pat in [r'第[一二三四五六七八九十\d]+次', r'首次', r'起兵', r'之乱', r'叛乱', r'政变', r'北伐', r'南渡', r'推行', r'提出', r'继位', r'称帝', r'建国', r'建立']:
        s = re.sub(pat, '', s)
    return re.sub(r'\s+', '', s).strip()


def _parse_year_range(stage_years):
    """从阶段年代范围字符串里解析出 (起年, 止年)，解析失败返回 (None, None)。"""
    nums = re.findall(r'\d{3,4}', stage_years)
    if len(nums) >= 2:
        return int(nums[0]), int(nums[1])
    if len(nums) == 1:
        return int(nums[0]), int(nums[0])
    return None, None


def dedupe_stage_outline(items, prior_stage_events, stage_years):
    """代码层硬去重 + 年代越界剔除（兜底，不靠 AI）。
    items: 本阶段大纲 [(event, point, year_start), ...]
    prior_stage_events: 前面阶段已列事件集合（已归一化的事件名 + 年份），用于跨阶段去重
    stage_years: 本阶段年代范围，用于剔除越界事件。
    返回去重后的 items。"""
    y_start, y_end = _parse_year_range(stage_years)
    kept = []
    for event, point, year in items:
        # 1) 越界剔除：起始年份不在本阶段范围内（仅在有有效范围时启用）
        if y_start is not None and y_end is not None:
            if year < y_start or year > y_end:
                continue
        # 2) 跨阶段去重：与前面阶段已列事件重名（归一化后）或同年同核心词
        norm = _normalize_event_name(event)
        dup = False
        for (pn, py) in prior_stage_events:
            if norm and norm == _normalize_event_name(pn):
                dup = True
                break
            if year == py and norm and (norm in _normalize_event_name(pn) or _normalize_event_name(pn) in norm):
                dup = True
                break
        if dup:
            continue
        kept.append((event, point, year))
    return kept


def review_all_outlines(topic, overview, stages_with_outlines):
    """让 AI 对全部阶段的大纲做一次性全局校验，返回修正后的大纲。
    stages_with_outlines: [(stage_no, stage_name, stage_years, stage_summary, items), ...]
    返回同结构列表，items 为校验修正后的 [(event, point, year_start), ...]。
    校验项：跨阶段重复、越界、史实遗漏、同年过碎、事件名不规范、笔误。
    仍以代码层 dedupe_stage_outline 兜底（AI 校验可能漏），故此处不做代码去重。"""
    # 拼成 AI 可读的全文大纲
    outline_text = ""
    for stage_no, stage_name, stage_years, stage_summary, items in stages_with_outlines:
        outline_text += f"\n# 第{stage_no}阶段｜{stage_name}（{stage_years}）｜{stage_summary}\n"
        for event, point, year_start in items:
            outline_text += f"{event}｜{year_start}｜{point}\n"

    prompt = f"""
你是严谨的历史学家，正在为【{topic}】做科普大纲的【全局校验】。下面是已生成的全部阶段大纲，每个事件一行"事件名｜起始年份｜要点概述"。

【{topic}】背景总览（事实基线）：
{overview}

【全部阶段大纲】：
{outline_text}

请做一次性全局校验，修正以下问题：
1. 【跨阶段重复】同一事件在多个阶段出现的，只在它【起始年份所属的阶段】保留一条，其余删除。
2. 【年代越界】事件起始年份不在所属阶段年代范围内的，移到正确阶段；找不到合适阶段的删掉。
3. 【史实遗漏】本朝代重大史实（关键战役、重要人物登场/去世、重大制度、政权更替）遗漏的，补到对应阶段正确位置。
4. 【同年过碎】同一年发生、本属同一件事的几个环节被拆成多条的，合并成一条（事件名取概括名，要点写清各环节）。
5. 【事件名不规范】事件名不得带括号年代（如"桓温专权（345-372）"应改为"桓温接掌荆州军权"）；事件名要具体、指明是什么事，不要用阶段名当事件名。
6. 【笔误】人名/地名/战役名笔误的，修正（如"泗水"应为"淝水"）。

输出要求：
- 严格保持原阶段数和阶段顺序，逐阶段输出修正后的大纲。
- 每个阶段先输出一行标题"第N阶段｜阶段名（年代）"，紧跟该阶段的事件列表。
- 每个事件一行"事件名｜起始年份｜要点概述"，行首不加序号，竖线｜分隔三段。
- 只输出修正后的大纲，不要解释、不要评论、不要输出背景总览。
- 若某阶段无需修改，原样输出该阶段。
"""
    print("\n校验全部大纲中（AI 全局自检自修）...")
    raw = call_ai(prompt)

    # 解析 AI 返回，按"第N阶段｜"标题分段
    reviewed = []
    cur = None  # (stage_no, stage_name, stage_years, stage_summary, items)
    for ln in raw.splitlines():
        ln = ln.strip()
        if not ln:
            continue
        # 标题行：第N阶段｜阶段名（年代）
        m = re.match(r'^第(\d+)阶段[｜|](.+?)（(.+?)）', ln)
        if m:
            if cur is not None:
                reviewed.append(cur)
            sno = int(m.group(1))
            sname = m.group(2).strip()
            syears = m.group(3).strip()
            # 从原 stages_with_outlines 取回 summary（AI 不输出 summary）
            summary = ""
            for o in stages_with_outlines:
                if o[0] == sno:
                    summary = o[3]
                    break
            cur = (sno, sname, syears, summary, [])
            continue
        # 事件行：事件名｜年份｜要点
        if cur is not None and '｜' in ln:
            parts = ln.split('｜', 2)
            if len(parts) == 3:
                event, year_str, point = [p.strip() for p in parts]
                m2 = re.search(r'(\d{3,4})', year_str)
                if event and m2:
                    cur[4].append((event, point, int(m2.group(1))))
    if cur is not None:
        reviewed.append(cur)

    # 兜底：若 AI 返回解析失败（没解析出任何阶段），原样返回原大纲
    if not reviewed or all(len(r[4]) == 0 for r in reviewed):
        print("  ⚠️ AI 校验返回解析失败，沿用原大纲。")
        return stages_with_outlines
    return reviewed


# ---- 第四步：单条大纲扩写（沿用已验证约束）----
def generate_section_details(topic, stage_name, stage_years, event, point, event_year, prev_detail, prior_events, is_stage_first=False):
    """针对一个历史事件，写一段详细科普正文（讲清起因/经过/结果/脉络推进，集中讲透，不重复）。
    stage_years: 本阶段的年代范围（如"317-322"），开篇定调时点明年代用。
    event_year: 本事件的起始年份，用于约束"不提及时间在本事件之后的事件"。
    prev_detail: 上一条已生成正文的尾部，用于自然衔接（第一条为空字符串）。
    prior_events: 本阶段在本次之前已扩写完的事件列表 [(事件名, 要点, 年份), ...]，用于避免背景重讲。
    is_stage_first: 本条是否是本阶段（也是本文件）的第一个事件。True 时走"开篇定调"分支——
        先一两句定调本阶段年代/主题，再进本事件，不复述上一阶段结尾，避免文件1结尾≈文件2开头的重叠。"""
    prev_block = ""
    if is_stage_first:
        # 跨阶段（跨文件）首条：开篇定调，不复述上一阶段结尾，避免文件衔接处重叠
        prev_block = f"""
【开篇定调——本条是本阶段/本文件的第一条，重要】
本条不是串行下一条，而是新阶段的开篇。先简短开篇（两三句即可）：点明本阶段【{stage_name}（{stage_years}）】处在该朝代的什么年代、什么主题、上承什么局面（一句话带过，不复述上一阶段结尾的具体史实/人物/经过），然后直接进入本事件【{event}（{event_year}年）】。
【严禁与上一文件结尾重叠——重要】不要复述上一阶段最后讲过的内容（人物、事件经过、结论），上一阶段讲过的默认读者已知道。本文件的开篇要让读者一眼看出"进入新阶段了"，而不是和上个文件的结尾接得像同一段。只点一句"上阶段讲到某局面之后"即可，不展开。
"""
    if prev_detail or prior_events:
        prev_block += f"""

【上下文——重要】
"""
        if prior_events:
            prior_list = "\n".join(f"- {e}（{y}年）｜{p}" for e, p, y in prior_events)
            prev_block += f"""
【本阶段已经讲过的事件】（这些事件时间都在本次事件之前或同期，经过/背景/人物已讲过，你不要再重讲；
讲到关联时一句带过即可，如"前面讲过的王敦之乱平定后……"，然后直接进入本次事件的新内容）：
{prior_list}
"""
        if prev_detail:
            prev_block += f"""
【上一条大纲的正文结尾】（这是上一条讲到的地方，仅供你知道故事进行到哪了）：
{prev_detail}
"""
        prev_block += f"""
【衔接要求——重要】
你的本条要从上一条讲到的地方之后接着讲，沿时间线往前推进，不要往回看。只做简短的时间过渡（如"第二年""就在这时""同一年的冬天""就在朝堂还在为这件事吵的时候"），一句话带过，直接进入本条事件的新内容。
严禁前情回顾：不要复述已讲过事件的经过、起因或结果；不要用"上回说到""刚才讲到""前文提到""我们之前说过""上一条说到""前面讲过的……之后""上一条讲到的地方之后"之类往回看的话。已讲过的部分，默认读者已经知道，不重讲。【已讲制度不展开】若【本阶段已经讲过的事件】里出现过某项制度/政策（如侨寄法、侨置郡县、土断），本条只能直接用名字一句带过它在本事件里的作用，不要再解释它是什么、怎么运作、有哪几条规定、有什么利弊——这些在它首次出现的那一条已经讲透，重讲就是重复。
【时间纪律——重要】本次事件发生在 {event_year} 年。只讲这一事件及其直接前因后果，不要提及时间在本事件【之后】的事件（那些是后面要讲的，不能提前剧透或关联）；只能引用时间在本事件【之前或同期】的已讲事件作背景，且引用时一句带过即可，不要展开回顾。
【严禁未来剧透】无论出于什么理由（铺垫影响、说明意义、吊胃口），都不得展开讲述时间在本事件【之后】才发生的事件的细节——不写其人物、年份、数字、经过，也不复述其后才出现的名言与场景。若必须点一句后续走向，只能用"后来……"这类不带任何细节的泛指一笔带过，把具体史实留给真正讲到那件事的条目。
【已讲名言/典故/典礼动作不重讲】凡在【本阶段已经讲过的事件】里已展开过的史书名言、成语典故、典礼上的标志性动作（如"太阳下同万物，苍生何由仰照""闻鸡起舞""拉王导同坐"），本条提到时只用名字一句带过，不复述原话、不重描场景。
"""
    prompt = f"""
请针对【{topic}】的【{stage_name}】阶段中的这个历史事件，写一篇非常详细的科普文章，面向完全不懂历史的普通读者。

事件：{event}
要点概述：{point}

要求：
1. 字数不少于 800 字。
2. 把这个事件讲透：必须含具体年份、人物、经过；
   要讲清它的起因（为什么会发生）、经过（怎么发展的）、结果（最后怎样），
   结果之后再用一两句话点明"这件事把朝代推进到了什么状态、引向了下一步什么局面"——这是本朝代的【脉络推进】，不是宏观影响评述。
   让小白听完能完整理解这件事的前因后果，并看清它在朝代发展链条上的位置。
3. 若该事件的背景在前面条目已交代过，本条一句话带过或不提，不要把已讲过的内容当背景再重讲一遍。
4. 拒绝宏观叙事，不要泛泛而谈，只要具体史实细节。
5. 直接输出正文内容，使用 Markdown 格式排版（用 ### 作为标题）。
6. 语言口语化、娓娓道来，像跟朋友讲故事，不要学术腔、不要书面语、不要生硬罗列。
7. 与上一条的衔接只做简短的时间/场景过渡，直接进入新内容，严禁前情回顾和复述（详见下方【衔接要求】）。
8. 第一次提到的人名、官职、地名、制度时，顺带用大白话解释一句是什么、干什么的、在哪儿，
   不要默认读者知道（例如"刺史"是地方军政长官、"建康"就是今天的南京、"内阁"是皇帝的秘书班子）。
   【制度不重复解释——重要】某项制度/政策（如侨寄法、侨置郡县、土断）若在【本阶段已经讲过的事件】里已出现过，
   本条提到它时只直接用名字、一句带过其在本事件中扮演的角色即可，【绝对不要】再解释它"是什么、怎么运作、有哪几条规定"。
   制度的定义和机制只在它首次出现的那一条讲透，后续条目一律不重讲机制、不重列其规定与利弊。
9. 【最高原则——严格忠于史实，绝不揣测编造】整篇内容必须基于真实历史，只能写有史料依据的事实。
   不得虚构任何内容，包括但不限于：人物对话、心理活动、细节动作、场景描写、人物结局、事件因果。
   对话不要凭空编写（不要写"某某说：……"除非确有史料记载原话）；人物心理不要揣测（不要写"他心里一沉""他暗想"）；
   场景细节不要脑补（不要写"风很大""烛火摇曳"等无依据描写）。
   拿不准的细节宁可略去或用"据记载""据史书"等留白表述，绝不臆造。
   例如某人是被杀、自杀、逃亡还是病故，必须区分清楚，不能图省事一律写成"被杀"。
   【年份/履历必须核实——重要】具体年份（如某皇帝被俘、某城被破、某战役发生）必须准确，不得前后矛盾、不得张冠李戴；
   人物履历（如"跟着某皇帝打天下"）必须有据，不得夸大或编造——拿不准的履历宁可不写，绝不杜撰"跟着某开国皇帝打天下"之类无依据说法。
10. 第一次提到次要人物时，只交代其确凿的身份与立场，不展开未经核对其后人生结局；
    若确知其结局，简述即可，不添油加醋。
11. 【只讲本事件的新内容——重要】本条要讲透的是【本事件本身】的起因/经过/结果/脉络推进。
    凡在【本阶段已经讲过的事件】里已讲过的制度机制、人物背景、前因铺垫、名言典故、典礼上的标志性动作，本条不得再展开；
    只能引用其名做一句过渡（如"沿用前面讲过的侨寄法安置这批新流民"），把篇幅留给本事件真正新增的史实与脉络推进。
    【名言/典故/典礼动作只讲一次】史书名言（如"若太阳下同万物，苍生何由仰照"）、成语典故（如"闻鸡起舞""中流击楫"）、典礼上的标志性动作（如"拉王导同坐龙椅"），只在它【真正发生的那一条】展开讲；
    本条若不是该名言/典故/动作的归属事件，提到时只用名字一句带过，不得复述原话、不得重描场景。
    【重复出场人物只首次详介——重要】同一人物（如王敦、王导）若在前面条目已详细介绍过其身份、家族、官职、兵力、驻地，本条再次提到时【只用人名】，不得再重复介绍"他是某某堂兄、手握某地重兵、驻守某处"这些已知信息；只在人物【首次登场】的那一条详介身份，后续一律一句带过或直接用人名推进叙事。
12. 【影响一律不写——重要】本条【不得单独写"影响"段】。前面条目已经反复出现"皇权孱弱""门阀专权""谁有兵谁说了算""为后世埋伏笔"之类结论性影响，重复读起来啰嗦。
    本条只在交代完"结果"后，用一两句话点明这件事【把朝代推进到什么状态、引向下一步什么局面】即可，不要写宏观影响评述，不要总结历史意义。
    例外：本事件若是本朝代重大转折（如开国、迁都、制度确立、关键战役定鼎、亡国），可多写一两句它独有、挪不走的改变；但同一结论前面已讲过的，不得再重复。
    【脉络推进不重复宏观论断——重要】"脉络推进"句只写【本事件独有】的推进（这件事把朝代从状态A推到状态B、引向下一步C），不得重复前面条目已写过的宏观结论（如"皇权孱弱""门阀专权""谁拳头硬谁老大""工具人皇帝"）。判断标准：把这句挪到前面任一已讲事件里都说得通 → 删掉，只留挪不走的。
{prev_block}
"""
    print(f"  -> 正在扩写: {event}")
    return call_ai(prompt, max_tokens=12000)


def get_tail(text, n=400):
    """取正文结尾 n 字，作为下一条衔接的上下文。"""
    text = text.strip()
    return text[-n:] if len(text) > n else text


# ---- 第五步：转 Word ----
def markdown_to_word(md_content, file_label):
    """把一个阶段的扩写合并为一个 Word。file_label 如 '宋朝_1_开国奠基'。"""
    print(f"  生成 Word: {file_label}...")
    word_file = f"{file_label}.docx"
    doc = Document()

    # 设置正文中文字体（默认字体不含中文会显示异常）
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(12)

    # 大标题：file_label 中的下划线换成中圆点展示，居中 + 加大字号
    title = doc.add_heading(file_label.replace('_', ' · '), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)

    # 按 Markdown 的分隔线切分章节，逐块解析
    sections = md_content.split('\n---\n')
    for sec in sections:
        for line in sec.splitlines():
            line = line.rstrip()
            if not line:
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            else:
                para = doc.add_paragraph()
                add_runs_with_bold(para, line)
        if sec.strip():
            doc.add_paragraph()  # 章节间空行

    doc.save(word_file)
    print(f"  🎉 已生成: {word_file}")


def review_section_text(topic, stage_name, md_content):
    """让 AI 对一个阶段扩写好的正文做校验自修，返回修正后的 Markdown 正文。
    校验项：与同阶段前文重复（典故/人物介绍/宏观结论）、史实错误（年份/履历/亲属关系/生死时间/驻防）、
    时间错位、前情回顾、结尾截断。返回修正版正文；解析失败则原样返回（兜底）。"""
    prompt = f"""
你是严谨的历史学家，正在为【{topic}】的【{stage_name}】阶段做科普正文的【校验自修】。
下面是该阶段已扩写好的全文（Markdown 格式，用 ### 作小标题，多个事件用 --- 分隔）。

【正文】：
{md_content}

请做一次性校验自修，修正以下问题：
1. 【重复】同一史书名言/典故/典礼动作（如"拉王导同坐龙椅""太阳下同万物苍生何由仰照""闻鸡起舞""中流击楫"）在多个事件里展开讲的，只在它【真正归属的那个事件】保留完整讲述，其余事件里改为一句话带过（用名字指代，不复述原话、不重描场景）。
2. 【重复】同一人物（如王敦、王导）的身份/家族/官职/兵力/驻地，在多个事件里反复介绍的，只在【首次登场】的事件保留完整介绍，其余事件里删掉重复介绍，直接用人名推进叙事。
3. 【重复】同一宏观结论（如"皇权孱弱""门阀专权""谁拳头硬谁老大""工具人皇帝""为后世埋下伏笔""开先例"）在多个事件末尾重复出现的，只在最贴切的一个事件里保留一句，其余删掉。
4. 【史实错误·年份履历】具体年份（如某皇帝被俘、某城被破、某战役）、人物履历（如"跟着某皇帝打天下"）、事件因果，若有错误或前后矛盾，据史修正；拿不准的宁可用"据记载"留白或略去，绝不杜撰履历。
5. 【史实错误·亲属关系——重要】人物之间的父子/兄弟/叔侄等亲属关系必须准确。如王含是王敦之兄，不是王导之父；王导之父是王裁，王敦之父是王基。不得把人物辈分、亲缘关系搞乱。拿不准的关系宁可不提，绝不杜撰。
6. 【史实错误·人物生死时间——重要】不得让已故人物登场行动。如周访320年已死，322年王敦起兵时不可能被"调来驻防"；必须核实人物在该事件发生时是否在世，已故者不得出现参与动作，只能作回溯提及。
7. 【史实错误·驻防地点】人物的镇守地、官职驻地必须准确（如刘隗322年以镇北将军都督淮阴，非合肥；戴渊镇合肥）。不得张冠李戴把甲的驻地写成乙。拿不准的宁可不写具体地点。
8. 【时间错位】把发生在某年的典故/事件，错放在另一年的叙事框架下讲的，调整时间表述，让时间线自洽。
9. 【前情回顾】出现"上回说到""刚才讲到""前文提到""前面讲过的……之后"等往回看的话，删掉，直接进入新内容。
10. 【结尾截断——重要】检查正文最后一段是否话说一半戛然而止（如以"最过分的是""但是""然而"等开头却没下文）。若发现截断，把该句补全成完整段落，让正文自然收尾；不要新增事件，只把最后那段补完。
11. 【内部一致性·人物驻地/官职——重要】同一人物的镇守地、官职、身份，在全文多处提到时必须【完全一致】，不得一处写甲、另一处写乙。如戴渊镇守地，全文要么都写合肥、要么都写芜湖，不得前文合肥后文芜湖；刘隗镇守地全文统一。发现不一致的，据史实统一为正确地点，并把所有提及处改成同一表述。
12. 【内部一致性·关键节点年份——重要】同一关键节点（如某皇帝称帝/称王、某城被破、某战役爆发）的年份，在全文多处提到时必须【完全一致】，不得一处写A年、另一处写B年。如司马睿317年称晋王、318年称帝，全文凡是提到"称帝"年份的都必须是318，凡是提到"称晋王/称王"的都必须是317，不得把"317年称帝"这种口误与"318年称帝"并存。发现不一致或口误的，统一为正确年份。
13. 【年份精度·重大任命】重大人事任命（如皇帝任命某将镇某地）的年份应据史实核准。如司马睿任命戴渊、刘隗外出镇守是321年（王敦322年起兵前一年），不得提前到320年。拿不准的宁可不写具体年份，用"不久""第二年"等相对时间表述。

输出要求：
- 保持原有的事件数量、顺序、### 小标题、--- 分隔线结构不变。
- 直接输出修正后的完整正文（Markdown 格式），不要解释、不要评论、不要列修改清单。
- 只改问题处，没问题的段落原样保留，不要重写正确的内容。
- 务必输出完整全文，不要省略、不要用"……（后略）"代替。
"""
    print("  校验正文（AI 自检自修）...")
    revised = call_ai(prompt, max_tokens=12000)
    # 兜底：若返回明显过短或无 ### 标题，视为解析失败，沿用原文
    if not revised or len(revised) < len(md_content) * 0.5 or '###' not in revised:
        print("  ⚠️ AI 正文校验返回异常，沿用原文。")
        return md_content
    return revised


def safe_filename(name):
    """把阶段名里不能当文件名的字符替换掉。"""
    return re.sub(r'[\\/:*?"<>|]', '_', name)


if __name__ == "__main__":
    # 1. 背景总览（1次调用，全局事实基线）
    overview = generate_overview(TOPIC)

    # 2. 兴衰阶段划分（1次调用，通吃所有朝代）
    stages = generate_stages(TOPIC, overview)
    if not stages:
        print("\n⚠️ 阶段划分失败，请检查 API 或重试。")
        raise SystemExit(1)

    print(f"\n【{TOPIC}】共划分为 {len(stages)} 个兴衰阶段：")
    print("-" * 60)
    for no, name, years, summary in stages:
        print(f"  阶段{no}｜{name}｜{years}｜{summary}")
    print("-" * 60)

    # 3. 先把【所有阶段的大纲】全部梳理出来，再去按大纲扩写写文件
    #    这样跨阶段去重能在"全部大纲生成完"后一次性、全局地做，重复/越界事件在写正文前就被剔除，
    #    避免边生成边写导致前面文件已落盘、后面才发现重复来不及改。
    print(f"\n第三步：一次性梳理全部 {len(stages)} 个阶段的大纲...")
    prior_stage_events = []  # 前面阶段已列事件（事件名+年份），跨阶段累积，用于大纲去重
    raw_outlines = []  # [(stage_no, stage_name, stage_years, stage_summary, items), ...]
    for stage_no, stage_name, stage_years, stage_summary in stages:
        print(f"\n  生成【阶段{stage_no}：{stage_name}（{stage_years}）】大纲...")
        items = generate_stage_outline(TOPIC, overview, stage_no, stage_name, stage_years, stage_summary, prior_stage_events=prior_stage_events)
        # 代码层兜底：剔除越界事件 + 跨阶段重名/同年同核事件
        items = dedupe_stage_outline(items, prior_stage_events, stage_years)
        print(f"  本阶段大纲共 {len(items)} 个事件（按起始年份排序）：")
        for idx, (event, point, year_start) in enumerate(items, 1):
            print(f"    {idx}. [{year_start}年] {event}｜{point[:40]}")
        raw_outlines.append((stage_no, stage_name, stage_years, stage_summary, items))
        # 累积本阶段大纲事件，供后续阶段大纲去重（以"大纲层已列"为准，先到先得）
        for event, point, year_start in items:
            prior_stage_events.append((event, year_start))

    # 3.5 AI 全局校验全部大纲：自检重复/越界/遗漏/同年过碎/命名/笔误，返回修正版
    all_outlines = review_all_outlines(TOPIC, overview, raw_outlines)
    print(f"\n校验后大纲（共 {len(all_outlines)} 个阶段）：")
    print("-" * 60)
    for stage_no, stage_name, stage_years, stage_summary, items in all_outlines:
        print(f"  阶段{stage_no}｜{stage_name}（{stage_years}）｜{len(items)}个事件")
        for idx, (event, point, year_start) in enumerate(items, 1):
            print(f"    {idx}. [{year_start}年] {event}｜{point[:40]}")
    print("-" * 60)

    print(f"\n{'=' * 60}")
    print(f"全部大纲梳理+校验完成。接下来按大纲逐阶段扩写、落盘 Word。")
    print(f"{'=' * 60}")

    # 4. 按已梳理好的大纲，逐阶段串行扩写 → 落盘一个 Word
    #    跨阶段也保持衔接（prev_tail 跨阶段不断），让所有 Word 读起来是一条连贯的历史
    prev_tail = ""
    for stage_no, stage_name, stage_years, stage_summary, items in all_outlines:
        print(f"\n{'=' * 60}")
        print(f"阶段 {stage_no}/{len(stages)}：{stage_name}（{stage_years}）")
        print(f"{'=' * 60}")

        if not items:
            print("  ⚠️ 本阶段大纲为空（可能因去重/越界剔除），跳过。")
            continue

        # 4.1 串行扩写本阶段每个事件
        #    prior_events 累积本阶段已扩写完的事件（事件名+要点+年份），喂给下一条避免背景重讲
        #    跨阶段首条走"开篇定调"分支（is_stage_first=True），避免文件衔接处重叠
        buffer = []
        prior_events = []  # 本阶段已讲过的事件，跨阶段不累积（每个新阶段重新开始）
        for idx, (event, point, year_start) in enumerate(items):
            is_first = (idx == 0)
            detail = generate_section_details(TOPIC, stage_name, stage_years, event, point, year_start, prev_tail, prior_events, is_stage_first=is_first)
            buffer.append(detail)
            prior_events.append((event, point, year_start))  # 记下已讲，供本阶段后续事件引用
            prev_tail = get_tail(detail)  # 跨阶段衔接不断
            print(f"    ✔ [{year_start}年] {event[:30]}... {len(detail)}字")

        # 4.2 本阶段扩写完，先 AI 校验正文（自检重复/史实/时间错位/前情回顾），再落盘 Word
        chunk_md = "\n\n---\n\n".join(buffer)
        chunk_md = review_section_text(TOPIC, stage_name, chunk_md)
        topic_short = TOPIC.replace("历史", "").strip() or TOPIC
        file_label = f"{topic_short}_{stage_no}_{safe_filename(stage_name)}"
        markdown_to_word(chunk_md, file_label)

    print(f"\n全流程结束！请去当前文件夹查看 {len(stages)} 个 Word 文件。")
