# -*- coding: utf-8 -*-
"""历史人物人生智慧 · 单期播客口播稿生成器

流程：背景卡（1次调用）→ 5段脚本串行（5次调用，每段衔接前一段）→ 拼接 → Word。
5段结构固定，对应策划案里的【开场钩子/历史故事/决策拆解/当代映射/结尾金句+互动】。
"""
import re
from openai import OpenAI
import httpx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ================= 配置区域 =================
API_KEY = "sk-Qaes2KCUDHr67GZoif13ySqsuFsD7NYWAnLoVcNcAlv3mcXW"
MODEL_NAME = "DeepSeek-V4-Flash"
#MODEL_NAME = "GLM-5.2-internal"
# 本期选题：人物 + 主题标签（当代困境映射）
# ================= 配置区域 =================
EPISODE_TITLE = "苏轼黄州——35岁被裁员的自救指南"
EPISODE_PERSON = "苏轼"
EPISODE_THEME = "失业/被裁员的绝境重建"
# ============================================

client = OpenAI(
    api_key=API_KEY,
    base_url="http://aiserver.hisi.huawei.com/v1",
    # 绕过 Windows 注册表里的 IE 系统代理（华为 proxycn2:8080），否则 504
    http_client=httpx.Client(trust_env=False)
)

_BOLD_RE = re.compile(r'\*\*(.+?)\*\*')


def add_runs_with_bold(paragraph, text):
    """把含 **加粗** 的文本拆成多个 run，加粗部分 bold=True；未配对的 ** 直接清除。"""
    if text.count('**') >= 2:
        pos = 0
        for m in _BOLD_RE.finditer(text):
            if m.start() > pos:
                paragraph.add_run(text[pos:m.start()])
            run = paragraph.add_run(m.group(1))
            run.bold = True
            pos = m.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])
    else:
        paragraph.add_run(text.replace('**', ''))


def call_ai(prompt):
    """调用大模型的通用函数"""
    print("  思考中...", end="", flush=True)
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": prompt}]
    )
    print("完成!", flush=True)
    return response.choices[0].message.content


# ---- 第一步：背景卡（给5段脚本统一喂入的事实基线，避免各段事实打架）----
def generate_context_card(person, theme):
    prompt = f"""
你是历史顾问。请为播客脚本准备【{person}】的背景卡，主题方向是【{theme}】。
只输出事实，不要议论，按下面6个小标题逐段输出，每段2-4句话，必须含具体年份/地名/人名：
### 核心困境
### 时间地点
### 关键人物
### 转折事件
### 关键决策（3条）
### 历史结果与影响
"""
    print(f"\n第一步：生成【{person}】背景卡...")
    return call_ai(prompt)


# ---- 第二步：5段脚本 ----
# 每段：固定结构 + 字数 + 风格约束 + 喂入背景卡保证事实一致
# 串行执行：每段都把前面已生成的段落作为上下文喂入，保证衔接自然流畅
SECTIONS = [
    # (段名, 字数, 段落指令)
    ("开场钩子", 600,
     "先抛出一个当代人的真实困境（与主题相关），用第二人称'你'拉近距离，"
     "制造悬念，结尾用一句话过渡到本期人物。不要学术腔，像跟朋友聊天。"),
    ("历史故事", 3500,
     "讲{person}在类似困境中的真实经历，要有细节、画面感、情绪起伏，按时间顺序展开。"
     "必须使用背景卡里的事实，不得虚构年份/人物/事件。不要议论，先把故事讲好。"
     "要讲清这件事的来龙去脉——当时是什么局势、这件事为什么会发生、牵涉到哪些人，"
     "让完全不了解这段历史的听众也能听懂前因后果；第一次提到的人名/官职/地名顺带交代一句。"
     "为讲清来龙去脉，可适当超出字数。"
     "衔接要求：顺着开场钩子结尾的引入，直接从故事的起因开始按时间顺序往下讲，"
     "不要用'刚才说到''咱们接上''好，咱们接着聊'这类回溯套语，"
     "不要复述开场钩子里已经提过的'结果'（如被贬到某地），开场钩子是悬念预告，不是已讲过的情节。"),
    ("决策拆解", 2200,
     "提炼{person}做了哪几个关键选择（3-5个），每个选择讲清背后逻辑，用'第一性原理'思维，"
     "不要鸡汤。每个选择用'### 选择N：xxx'作为小标题。"
     "开篇要自然承接【历史故事】段的结尾，不要重复历史故事的内容，直接进入决策分析。"),
    ("当代映射", 1800,
     "把{person}的决策翻译成今天的行动指南，给出具体、可执行的步骤，不要泛泛而谈。"
     "可以用'如果你今天……那么……'的句式。"
     "开篇要用一句话把上面的历史决策过渡到今天，让听众意识到'古人怎么选'和'今天怎么选'的连接。"),
    ("结尾金句与互动", 600,
     "先用一句总结性金句收束全篇，再抛出一个问题引导听众留言，最后预告一句下期方向。"
     "金句要有力量但不要烂俗。"
     "开头不要再说'最后'之类的硬转折，要让听众感到这是水到渠成的收束。"),
]


def build_section_prompt(name, word_count, instruction, person, theme, context_card, prev_sections):
    # prev_sections: 前面已生成段落的拼接文本（开场钩子段为空字符串）
    prev_block = ""
    if prev_sections:
        prev_block = f"""

【前面已生成的段落】（你的段落要自然承接它们，不要重复内容，不要硬转折）：
{prev_sections}
"""
    return f"""
你在为一档叫《古人怎么选》的历史播客写口播稿的【{name}】段。
本期人物：{person}；主题方向：{theme}。

【背景卡】（事实基线，必须遵守，不得与之矛盾）：
{context_card}
{prev_block}
写作要求：
1. 字数约 {word_count} 字（宁可少写也不要注水）。
2. {instruction}
3. 直接输出正文，用 Markdown 排版：本段大标题用 '## {name}'，内部小标题用 '### '。
4. 不要重复背景卡的原话，要把事实融进叙述里。
5. 不要写'本段完''字数'之类元信息。
6. 语言要口语化、娓娓道来，像跟朋友聊天讲故事，不要学术腔、不要书面语、不要生硬罗列要点。
7. 同一地名/人名在全文保持一致：第一次出现时给出古今对照（如'黄州，也就是今天的湖北黄冈'），后续统一用同一个称呼，不要一会儿古名一会儿今名让听众对不上号。
"""


def generate_all_sections(person, theme, context_card):
    print(f"\n第二步：串行生成 5 段脚本（每段衔接前一段）...")
    results = {}
    for name, wc, instr in SECTIONS:
        # 把前面已生成段落的文本按顺序拼起来，作为衔接上下文
        prev_text = "\n\n---\n\n".join(
            results.get(n, "") for n, _, _ in SECTIONS if n in results
        )
        prompt = build_section_prompt(name, wc, instr, person, theme, context_card, prev_text)
        results[name] = call_ai(prompt)
        print(f"  ✔ [{name}] {len(results[name])}字")
    # 按 SECTIONS 固定顺序拼接，段间用分隔线
    return "\n\n---\n\n".join(
        results.get(name, "") for name, _, _ in SECTIONS
    )


# ---- 第三步：转 Word ----
def markdown_to_word(md_content, episode_title):
    print("\n第三步：生成 Word 文档...")
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(12)

    title = doc.add_heading(episode_title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)

    sections = md_content.split('\n---\n')
    for sec in sections:
        for line in sec.splitlines():
            line = line.rstrip()
            if not line:
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            else:
                para = doc.add_paragraph()
                add_runs_with_bold(para, line)
        if sec.strip():
            doc.add_paragraph()

    word_file = f"{EPISODE_PERSON}_第1期口播稿.docx"
    doc.save(word_file)
    print(f"\n🎉 成功生成: {word_file}")


if __name__ == "__main__":
    # 1. 背景卡（串行，1次调用）
    context_card = generate_context_card(EPISODE_PERSON, EPISODE_THEME)

    # 2. 5段并发
    full_script = generate_all_sections(EPISODE_PERSON, EPISODE_THEME, context_card)

    # 3. 转 Word
    markdown_to_word(full_script, EPISODE_TITLE)
    print("\n全流程结束！请去当前文件夹查看你的 Word 文件。")
